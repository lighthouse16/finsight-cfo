import io
import re

from app.models.data_room import DataRoomParsePreview
from app.services.data_room.structured_parser import _records_from_rows, PARSEABLE_RECORD_KEYS, EXPECTED_FIELDS

try:
    import docx
except ImportError:
    docx = None

def parse_docx_bytes(record_key: str, file_bytes: bytes) -> DataRoomParsePreview:
    if docx is None:
        return DataRoomParsePreview(
            recordKey=record_key,
            statementType=PARSEABLE_RECORD_KEYS[record_key],
            parsedRecords=[],
            missingExpectedFields=sorted(EXPECTED_FIELDS[record_key]),
            unsupportedFields=[],
            rowCount=0,
            warnings=["DOCX parsing is unavailable because python-docx is not installed. Analysis was not updated."],
        )
    
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        text = ""
        for para in document.paragraphs:
            extracted = para.text
            if extracted:
                text += extracted + "\n"
        
        # Also attempt to parse tables
        for table in document.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text += row_text + "\n"
    except Exception as exc:
        return DataRoomParsePreview(
            recordKey=record_key,
            statementType=PARSEABLE_RECORD_KEYS[record_key],
            parsedRecords=[],
            missingExpectedFields=sorted(EXPECTED_FIELDS[record_key]),
            unsupportedFields=[],
            rowCount=0,
            warnings=[f"DOCX could not be parsed safely: {exc}. Analysis was not updated."],
        )
    
    warnings = []
    if not text.strip():
        warnings.append("DOCX contains no text or readable tables.")
        return DataRoomParsePreview(
            recordKey=record_key,
            statementType=PARSEABLE_RECORD_KEYS[record_key],
            parsedRecords=[],
            missingExpectedFields=sorted(EXPECTED_FIELDS[record_key]),
            unsupportedFields=[],
            rowCount=0,
            warnings=warnings,
        )
        
    rows = [["metric", "value"]]
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        parts = re.split(r'\s{2,}|\t', line)
        if len(parts) >= 2:
            rows.append([parts[0], parts[-1]])
        elif len(parts) == 1:
            match = re.search(r'^(.*?)\s+([\d.,()\-]+)$', line)
            if match:
                rows.append([match.group(1), match.group(2)])
            
    preview = _records_from_rows(record_key, rows)
    preview.warnings.extend(warnings)
    return preview
